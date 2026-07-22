"""Generate the system documentation as both .docx and .pdf from one source.

    python build_docs.py <output-dir>

The CONTENT list below is the single source of truth, so the two files cannot
drift apart.
"""

from __future__ import annotations

import sys
from pathlib import Path

# ── Document content ─────────────────────────────────────────────────────
# Block forms:
#   ("h1", text) ("h2", text) ("p", text) ("bullet", text)
#   ("code", [lines]) ("caption", text) ("pagebreak",)
#   ("table", [headers], [[row]], [col fractions])

TITLE = "A Multi-Agent Architecture for Clinical Document Assessment"
SUBTITLE = (
    "System Structure, Agent Roles, and Confidence-Weighted "
    "Inter-Agent Agreement"
)
DOCTYPE = "Technical System Documentation"
VERSION = "Version 0.3"
DATE = "July 2026"
DISCLAIMER = (
    "Research and educational use only. This system is not a medical device "
    "and produces no clinical advice."
)

CONTENT = [
    ("h1", "1. Abstract"),
    ("p",
     "This document describes the structure and operation of a multi-agent system for the "
     "assessment of clinical documents. Rather than posing a clinical question to a single "
     "language model, the system decomposes the task across a coordinated set of specialised "
     "agents: a supervisor that determines which clinical specialties are relevant, a "
     "dynamically sized swarm of specialist agents that analyse the source documents "
     "independently and in parallel, a judge that consolidates their reviews, and a safety "
     "agent that verifies every surviving claim against the original source text."),
    ("p",
     "A distinguishing feature of the architecture is that agent agreement is measured rather "
     "than assumed. Each specialist reports a self-assessed confidence, and the system computes "
     "a panel-level agreement score by averaging those confidences and discounting the result "
     "by their dispersion. This prevents a divided panel from presenting as a confident one, "
     "exposes dissenting specialists explicitly, and supplies the downstream judge with a "
     "quantitative signal about where the panel diverges."),
    ("p",
     "The system is implemented as a Python service exposing both a request/response API and a "
     "streaming event interface, with a browser client that renders each agent as it executes. "
     "This document sets out the architecture, the agent roles, the agreement methodology, the "
     "event protocol, the validation strategy, and the limitations of the approach."),

    ("h1", "2. Introduction and Motivation"),
    ("h2", "2.1 Problem statement"),
    ("p",
     "A clinical question posed against a patient record is rarely confined to one domain. A "
     "single record may simultaneously raise questions of renal function, glycaemic control, "
     "drug dosing, and cardiovascular risk. A single general-purpose model answering such a "
     "question must hold all of these perspectives at once, and its output offers no way to "
     "distinguish a conclusion that several lines of reasoning support from one that rests on a "
     "single thread."),
    ("p", "Two failure modes motivate the design of this system:"),
    ("bullet",
     "Ungrounded assertion. A language model may state a clinically plausible fact that does not "
     "appear in the supplied documents. Such a claim is indistinguishable in tone from a "
     "supported one."),
    ("bullet",
     "Undifferentiated confidence. A single model emits one answer with one implicit confidence. "
     "There is no way to tell whether that confidence reflects broad support or a narrow, "
     "contestable inference."),
    ("h2", "2.2 Design response"),
    ("p",
     "The architecture addresses the first failure mode with a dedicated verification stage that "
     "treats the source documents as the sole ground truth, and the second by replacing a single "
     "reasoner with a panel whose members answer independently and whose degree of agreement is "
     "computed explicitly. Neither mechanism makes the system clinically safe; both make its "
     "output easier to interrogate."),

    ("h1", "3. System Architecture"),
    ("h2", "3.1 Processing pipeline"),
    ("p",
     "Execution proceeds through four sequential stages. Stages 1, 3, and 4 each invoke a single "
     "agent; stage 2 invokes a variable number of agents concurrently. Agreement scoring is a "
     "deterministic computation performed between stages 2 and 3 and involves no model call."),
    ("code", [
        "     Clinical question  +  patient documents",
        "                       |",
        "                       v",
        "     +-------------------------------------+",
        "     |  STAGE 1   Supervisor Agent         |",
        "     |  selects specialties, names a lead  |",
        "     +-------------------------------------+",
        "                       |",
        "                       v",
        "     +-------------------------------------+",
        "     |  STAGE 2   Specialist Swarm         |",
        "     |  N agents, executed in parallel     |",
        "     |  each returns findings + confidence |",
        "     +-------------------------------------+",
        "                       |",
        "                       v",
        "     +-------------------------------------+",
        "     |  Agreement scoring  (deterministic) |",
        "     |  confidence averaging + dispersion  |",
        "     +-------------------------------------+",
        "                       |",
        "                       v",
        "     +-------------------------------------+",
        "     |  STAGE 3   Judge Agent              |",
        "     |  consolidates, resolves conflicts   |",
        "     +-------------------------------------+",
        "                       |",
        "                       v",
        "     +-------------------------------------+",
        "     |  STAGE 4   Safety Agent             |",
        "     |  verifies claims against documents  |",
        "     +-------------------------------------+",
        "                       |",
        "                       v",
        "               Verified report",
    ]),
    ("caption", "Figure 1. Four-stage processing pipeline with deterministic agreement scoring."),

    ("h2", "3.2 Agent roles"),
    ("p",
     "Every agent derives from a common base class that handles model invocation, optional "
     "token-level streaming, structured JSON parsing, and prompt assembly. Agents differ in "
     "their system prompt, their output schema, and the model assigned to them."),
    ("table",
     ["Agent", "Stage", "Responsibility", "Structured output"],
     [
         ["Supervisor", "1",
          "Reads the question and documents; determines which clinical specialties are required "
          "and designates one as lead.",
          "specialties[], lead_specialist, rationale"],
         ["Specialist", "2",
          "Instantiated once per specialty. Analyses the documents from a single domain "
          "perspective, quoting supporting evidence.",
          "specialty, findings, recommendation, evidence_quotes[], confidence"],
         ["Judge", "3",
          "Consolidates all specialist reviews into one coherent report, resolving conflicts "
          "between them.",
          "final_summary, key_findings_by_specialty, consolidated_recommendations[]"],
         ["Safety", "4",
          "Verifies each claim in the judge report against the original documents and flags "
          "anything unsupported.",
          "verified_findings[], flagged_issues[], final_summary, overall_confidence, is_safe"],
     ],
     [0.14, 0.08, 0.42, 0.36]),
    ("caption", "Table 1. Agent responsibilities and output schemas."),

    ("h2", "3.3 Dynamic swarm construction"),
    ("p",
     "The number of specialist agents is not fixed. It is determined at run time by the "
     "supervisor and bounded by a configurable ceiling (MAX_SPECIALISTS, default 10) to prevent "
     "unbounded cost on an ambiguous input. A record raising a single domain question produces "
     "one specialist; a complex multi-morbidity record may produce several. Specialists are "
     "dispatched concurrently, so wall-clock latency is governed by the slowest agent rather "
     "than their sum."),

    ("h2", "3.4 Externalised prompts"),
    ("p",
     "Agent instructions are not embedded in source code. Each role has a directory under "
     "skills/ containing a Markdown instruction file and an optional references/ folder of "
     "supporting material. Both are loaded into the system prompt when the agent is constructed. "
     "Agent behaviour can therefore be revised without modifying Python, which keeps prompt "
     "engineering separable from application logic and legible to non-programmers."),

    ("pagebreak",),
    ("h1", "4. Confidence-Weighted Inter-Agent Agreement"),
    ("p",
     "This section describes the mechanism by which the system quantifies agreement across the "
     "specialist panel. It is the component that distinguishes this architecture from a "
     "conventional sequential agent pipeline."),

    ("h2", "4.1 Why a plain mean is insufficient"),
    ("p",
     "Each specialist reports a self-assessed confidence c in the interval [0, 1]. The naive "
     "aggregate is the arithmetic mean. This is inadequate because the mean is insensitive to "
     "the distribution that produced it. Consider two panels of two specialists each:"),
    ("table",
     ["Panel", "Reported confidences", "Arithmetic mean", "Interpretation"],
     [
         ["A", "0.50 and 0.50", "0.50",
          "Both specialists are moderately confident. The panel concurs."],
         ["B", "0.10 and 0.90", "0.50",
          "One specialist is nearly certain, the other nearly certain of the opposite. "
          "The panel is split."],
     ],
     [0.10, 0.24, 0.18, 0.48]),
    ("caption", "Table 2. Two panels with identical means but incompatible interpretations."),
    ("p",
     "The two panels are epistemically very different, yet the mean reports them identically. "
     "Presenting panel B as a 50% result conceals a genuine disagreement that a reader would "
     "want to see. The aggregation must therefore be sensitive to dispersion as well as to "
     "central tendency."),

    ("h2", "4.2 Definition"),
    ("p",
     "Let C = {c1, ..., cn} be the confidences of the n participating specialists. The agreement "
     "score A is defined as the mean confidence discounted by a normalised measure of "
     "dispersion:"),
    ("code", [
        "    mean        =  (1/n) * SUM ci",
        "",
        "    sigma       =  population standard deviation of C",
        "",
        "    dispersion  =  min( sigma / 0.5 ,  1 )",
        "",
        "    A           =  mean * ( 1 - dispersion )",
    ]),
    ("p",
     "The normalising constant 0.5 is the maximum standard deviation attainable by values "
     "bounded to [0, 1], which occurs when half the panel reports 0 and half reports 1. Dividing "
     "by it maps sigma onto [0, 1], so dispersion is a proportion rather than an unbounded "
     "quantity, and a maximally divided panel receives an agreement score of zero irrespective "
     "of how confident its individual members were."),
    ("p",
     "The population standard deviation is used in preference to the sample standard deviation "
     "because the panel is the entire population of interest, not a sample drawn from a larger "
     "one."),

    ("h2", "4.3 Behaviour"),
    ("p", "Applying the definition to the two motivating panels, and to two further cases:"),
    ("table",
     ["Confidences", "Mean", "sigma", "Dispersion", "Agreement A", "Level"],
     [
         ["0.50, 0.50", "0.500", "0.000", "0.00", "0.500", "moderate"],
         ["0.10, 0.90", "0.500", "0.400", "0.80", "0.100", "none"],
         ["0.90, 0.85", "0.875", "0.025", "0.05", "0.831", "strong"],
         ["0.70 (single agent)", "0.700", "0.000", "0.00", "0.700", "moderate"],
     ],
     [0.26, 0.13, 0.13, 0.17, 0.17, 0.14]),
    ("caption", "Table 3. Agreement scores for representative panels."),
    ("p",
     "The two panels that share a mean of 0.500 are now separated by a factor of five. A "
     "single-agent panel has no dispersion and therefore scores its own confidence unmodified, "
     "which is the correct degenerate case: with no second opinion there is nothing to agree or "
     "disagree with."),

    ("h2", "4.4 Abstention"),
    ("p",
     "A specialist that finds nothing relevant in its domain is instructed to report a "
     "confidence of zero. Such a response is treated as an abstention and excluded from both the "
     "mean and the dispersion. This distinction is material: a dermatologist finding no "
     "dermatological content in a cardiology record is not disagreeing with the cardiologist, "
     "and admitting that zero into the mean would depress the panel score for a reason unrelated "
     "to consensus. Abstentions are counted and reported separately so that they remain "
     "visible."),

    ("h2", "4.5 Outlier identification"),
    ("p",
     "A specialist whose confidence lies more than one standard deviation from the panel mean is "
     "labelled an outlier, subject to a minimum absolute deviation of 0.15. The floor is "
     "necessary because standard deviation is scale-relative: in a tightly clustered panel "
     "reporting 0.95, 0.85 and 0.85, the standard deviation is approximately 0.047, and the "
     "specialist at 0.95 sits more than one such deviation from the mean despite differing from "
     "its colleagues by a clinically negligible margin. Without the floor the system would label "
     "ordinary variation as dissent. The floor was introduced after this behaviour was observed "
     "in a live run of the system."),

    ("h2", "4.6 Interpretation bands"),
    ("p",
     "The continuous score is additionally reported as a qualitative band, for presentation "
     "purposes only; the underlying score is always retained."),
    ("table",
     ["Band", "Range of A", "Reading"],
     [
         ["strong", "A >= 0.75", "Specialists are confident and closely aligned."],
         ["moderate", "0.50 <= A < 0.75", "Broad alignment, or alignment at moderate confidence."],
         ["weak", "0.25 <= A < 0.50", "Meaningful divergence or generally low confidence."],
         ["none", "A < 0.25", "The panel is divided, or no specialist reported confidence."],
     ],
     [0.14, 0.24, 0.62]),
    ("caption", "Table 4. Qualitative agreement bands."),

    ("h2", "4.7 Use of the score"),
    ("p",
     "The agreement report is consumed at three points. It is emitted to the client as a "
     "discrete event so the interface can display it as soon as the swarm completes; it is "
     "injected into the judge prompt, with an instruction to state explicitly where the panel "
     "diverges rather than smoothing the disagreement away; and it is returned in the final API "
     "response for downstream analysis."),

    ("pagebreak",),
    ("h1", "5. Execution Model and Event Protocol"),
    ("p",
     "The system exposes two interfaces over the same pipeline. A request/response endpoint runs "
     "the pipeline to completion and returns the aggregate result. A streaming endpoint emits a "
     "typed event at every state transition, allowing a client to render the pipeline as it "
     "executes rather than after it finishes."),
    ("table",
     ["Event", "Emitted when"],
     [
         ["triage_thinking / triage_done",
          "The supervisor begins and completes specialty selection."],
         ["specialists_spawned",
          "The swarm has been constructed; carries the specialty list and lead."],
         ["specialist_thinking / specialist_done",
          "Once per specialist, on dispatch and on completion."],
         ["consensus_done",
          "Agreement scoring has been computed over the completed swarm."],
         ["judge_thinking / judge_done", "The judge begins and completes consolidation."],
         ["safety_thinking / safety_done",
          "The safety agent begins and completes verification."],
         ["agent_stream", "Per token, for any agent generating output."],
         ["pipeline_complete", "Terminal event carrying the full aggregate result."],
     ],
     [0.37, 0.63]),
    ("caption", "Table 5. Streaming event types."),
    ("p",
     "Because agreement scoring is deterministic and requires no model call, consensus_done is "
     "emitted immediately upon completion of the final specialist, with no perceptible delay "
     "between the swarm finishing and the agreement panel appearing."),

    ("h1", "6. Implementation"),
    ("h2", "6.1 Technology"),
    ("table",
     ["Layer", "Technology", "Role"],
     [
         ["Service", "Python 3.11+, FastAPI, Uvicorn",
          "HTTP and WebSocket interfaces, request validation."],
         ["Agents", "OpenAI API, asyncio",
          "Model invocation; concurrent dispatch of the swarm."],
         ["Contracts", "Pydantic", "Typed schemas for every agent output and API payload."],
         ["Configuration", "pydantic-settings",
          "Environment-driven settings with typed defaults."],
         ["Client", "React 19, Vite", "Live rendering of pipeline state."],
         ["Testing", "pytest, pytest-asyncio",
          "Unit and integration tests against mocked model responses."],
     ],
     [0.18, 0.32, 0.50]),
    ("caption", "Table 6. Implementation stack."),

    ("h2", "6.2 Repository structure"),
    ("code", [
        "agents/",
        "    base.py          shared model client, streaming, JSON parsing",
        "    supervisor.py    stage 1",
        "    specialist.py    stage 2, plus swarm factory",
        "    judge.py         stage 3",
        "    safety.py        stage 4",
        "orchestrator/",
        "    pipeline.py      four-stage coordinator",
        "    consensus.py     agreement scoring (section 4)",
        "    events.py        typed streaming events",
        "api/",
        "    main.py          REST and WebSocket endpoints",
        "    schemas.py       request and response contracts",
        "config/",
        "    settings.py      environment-driven configuration",
        "skills/             externalised agent prompts (section 3.4)",
        "frontend/           React client",
        "tests/              test suite",
    ]),

    ("h2", "6.3 Configuration"),
    ("p",
     "All operational parameters are supplied through environment variables and read into a "
     "typed settings object. Model selection is per-role, so the supervisor and safety agents "
     "can be assigned different models from the specialists. Transport verification, request "
     "timeout, swarm ceiling, and permitted client origins are likewise configurable."),

    ("h1", "7. Validation"),
    ("p",
     "The system is exercised at three levels. Agent-level tests confirm that each role parses "
     "well-formed model output correctly and degrades gracefully on malformed output. "
     "Pipeline-level tests run all four stages against scripted model responses and assert the "
     "aggregate result, including the computed agreement figures. Interface-level tests exercise "
     "the API endpoints and their schemas."),
    ("p",
     "The agreement component is tested independently against the cases in section 4, together "
     "with adversarial inputs: absent confidence fields, non-numeric values, values outside "
     "[0, 1], panels in which every member abstains, and the empty panel. In each case the "
     "scoring must return a well-formed report rather than raise. At the time of writing the "
     "suite comprises twenty tests, all passing."),
    ("p",
     "Automated tests are supplemented by end-to-end execution against the live model API, since "
     "mocked responses cannot demonstrate that the deployed system functions. The outlier floor "
     "described in section 4.5 was added as a direct result of such a run."),

    ("h1", "8. Limitations and Ethical Considerations"),
    ("h2", "8.1 Limitations of the method"),
    ("bullet",
     "Self-reported confidence is not calibrated. The agreement score aggregates what agents "
     "assert about their own certainty, which is a stated disposition rather than a measured "
     "probability. A uniformly overconfident panel yields a high agreement score."),
    ("bullet",
     "Agreement is not correctness. Independently reasoning agents drawing on a shared "
     "underlying model may share its errors, and concordance among them offers no protection "
     "against a systematic mistake. A high score indicates consistency, not validity."),
    ("bullet",
     "Confidence is scalar. The score measures dispersion in stated certainty, not semantic "
     "disagreement in the findings themselves. Two specialists may report identical confidence "
     "while advancing incompatible recommendations."),
    ("bullet",
     "Verification is bounded by the documents. The safety agent can only establish that a claim "
     "is or is not supported by the supplied text. It cannot detect a clinically incorrect "
     "statement that the documents happen to support, nor an omission."),
    ("h2", "8.2 Intended use"),
    ("p",
     "The system is a research and educational artefact. It is not a medical device, has not "
     "been clinically validated, and must not be used to inform patient care. Any material "
     "processed by it is transmitted to a third-party model provider, which precludes the use of "
     "identifiable patient data absent an appropriate legal and technical basis. The safety "
     "stage reduces the incidence of ungrounded assertion; it does not render the output "
     "reliable, and no stage of the pipeline substitutes for review by a qualified clinician."),

    ("h1", "9. Conclusion"),
    ("p",
     "The architecture described here decomposes clinical document assessment across four stages "
     "and a dynamically sized panel of specialist agents, and makes two properties of the result "
     "explicit that a single-model approach leaves implicit: whether each claim is grounded in "
     "the source documents, and how far the panel agreed in producing it."),
    ("p",
     "The agreement mechanism is deliberately simple. It requires no additional model call, is "
     "fully deterministic and therefore reproducible, and is legible enough that a reader can "
     "reconstruct the score by hand from the reported confidences. Its limitations follow from "
     "that simplicity: it aggregates stated rather than calibrated confidence, and it measures "
     "consistency rather than correctness. Within those bounds it supplies something a single "
     "opinion cannot, namely an explicit and inspectable signal of when the constituent analyses "
     "did not converge."),
]

ACCENT = (0x1F, 0x4E, 0x79)
MUTED = (0x59, 0x59, 0x59)
HEADER_BG = "1F4E79"
ZEBRA_BG = "F2F6FB"
CODE_BG = "F2F4F7"


def toc_entries():
    """Top-level sections, for the contents page."""
    return [text for kind, *rest in ((b[0], *b[1:]) for b in CONTENT)
            if kind == "h1" for text in [rest[0]]]


# ── DOCX ─────────────────────────────────────────────────────────────────
def build_docx(path: Path) -> None:
    import docx
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = docx.Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(1))
    usable = section.page_width - section.left_margin - section.right_margin

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.15

    def shade(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        el = OxmlElement("w:shd")
        el.set(qn("w:val"), "clear")
        el.set(qn("w:fill"), fill)
        tcPr.append(el)

    def para(text="", *, size=11, bold=False, italic=False, color=None,
             align=None, space_before=0, space_after=7, font=None, style=None):
        pr = doc.add_paragraph(style=style)
        run = pr.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if font:
            run.font.name = font
            run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        if color:
            run.font.color.rgb = RGBColor(*color)
        if align is not None:
            pr.alignment = align
        pr.paragraph_format.space_before = Pt(space_before)
        pr.paragraph_format.space_after = Pt(space_after)
        return pr

    # Page number in the footer
    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run()
    for el, attrs, text in (
        ("w:fldChar", {"w:fldCharType": "begin"}, None),
        ("w:instrText", {"xml:space": "preserve"}, "PAGE"),
        ("w:fldChar", {"w:fldCharType": "end"}, None),
    ):
        node = OxmlElement(el)
        for k, v in attrs.items():
            node.set(qn(k), v)
        if text:
            node.text = text
        run._r.append(node)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(*MUTED)

    # ── Title page ──
    for _ in range(6):
        para(space_after=0)
    para(TITLE, size=22, bold=True, color=ACCENT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    para(SUBTITLE, size=12, italic=True, color=MUTED,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=34)
    para(DOCTYPE, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)
    para(VERSION, size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)
    para(DATE, size=11, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=46)
    para(DISCLAIMER, size=9, italic=True, color=MUTED,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ── Contents ──
    para("Contents", size=15, bold=True, color=ACCENT, space_after=12)
    for entry in toc_entries():
        pr = para(entry, size=11, space_after=6)
        pr.paragraph_format.left_indent = Inches(0.2)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ── Body ──
    for block in CONTENT:
        kind = block[0]

        if kind == "h1":
            para(block[1], size=15, bold=True, color=ACCENT,
                 space_before=18, space_after=8)
        elif kind == "h2":
            para(block[1], size=12.5, bold=True, color=ACCENT,
                 space_before=13, space_after=6)
        elif kind == "p":
            para(block[1])
        elif kind == "bullet":
            pr = doc.add_paragraph(block[1], style="List Bullet")
            pr.paragraph_format.space_after = Pt(5)
            for r in pr.runs:
                r.font.size = Pt(11)
        elif kind == "caption":
            para(block[1], size=9, italic=True, color=MUTED,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
        elif kind == "pagebreak":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        elif kind == "code":
            for i, line in enumerate(block[1]):
                pr = para(line or " ", size=8.5, font="Consolas",
                          space_after=12 if i == len(block[1]) - 1 else 0)
                pr.paragraph_format.line_spacing = 1.0
                pr.paragraph_format.left_indent = Inches(0.18)
                pPr = pr._p.get_or_add_pPr()
                el = OxmlElement("w:shd")
                el.set(qn("w:val"), "clear")
                el.set(qn("w:fill"), CODE_BG)
                pPr.append(el)
        elif kind == "table":
            headers, rows, fracs = block[1], block[2], block[3]
            t = doc.add_table(rows=1, cols=len(headers))
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.autofit = False
            widths = [int(usable * f) for f in fracs]

            for i, htxt in enumerate(headers):
                cell = t.rows[0].cells[i]
                cell.text = ""
                run = cell.paragraphs[0].add_run(htxt)
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                cell.paragraphs[0].paragraph_format.space_after = Pt(2)
                shade(cell, HEADER_BG)

            for ri, row in enumerate(rows):
                cells = t.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = ""
                    run = cells[i].paragraphs[0].add_run(val)
                    run.font.size = Pt(9.5)
                    cells[i].paragraphs[0].paragraph_format.space_after = Pt(2)
                    if ri % 2:
                        shade(cells[i], ZEBRA_BG)

            # Width must be set on every cell, not just the column.
            for row in t.rows:
                for i, cell in enumerate(row.cells):
                    cell.width = widths[i]
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        else:
            raise ValueError(f"unknown block: {kind}")

    doc.save(str(path))
    print(f"wrote {path}")


# ── PDF ──────────────────────────────────────────────────────────────────
def build_pdf(path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        BaseDocTemplate, Frame, KeepTogether, ListFlowable, ListItem,
        NextPageTemplate, PageBreak, PageTemplate, Paragraph, Preformatted,
        Spacer, Table, TableStyle,
    )

    accent = colors.Color(*[c / 255 for c in ACCENT])
    muted = colors.Color(*[c / 255 for c in MUTED])
    content_w = letter[0] - 2 * inch

    ss = getSampleStyleSheet()
    body_st = ParagraphStyle(
        "Body", parent=ss["Normal"], fontName="Helvetica", fontSize=10.5,
        leading=15.5, spaceAfter=7, alignment=TA_JUSTIFY,
    )
    h1_st = ParagraphStyle(
        "H1", parent=body_st, fontName="Helvetica-Bold", fontSize=15,
        leading=19, textColor=accent, spaceBefore=17, spaceAfter=8,
        alignment=0, keepWithNext=1,
    )
    h2_st = ParagraphStyle(
        "H2", parent=body_st, fontName="Helvetica-Bold", fontSize=12,
        leading=16, textColor=accent, spaceBefore=12, spaceAfter=5,
        alignment=0, keepWithNext=1,
    )
    cap_st = ParagraphStyle(
        "Cap", parent=body_st, fontSize=8.5, leading=12, textColor=muted,
        alignment=TA_CENTER, fontName="Helvetica-Oblique", spaceBefore=5,
        spaceAfter=12,
    )
    code_st = ParagraphStyle(
        "Code", parent=ss["Code"], fontName="Courier", fontSize=7.6,
        leading=9.6, textColor=colors.black, backColor=colors.HexColor("#" + CODE_BG),
        borderPadding=7, leftIndent=8, spaceAfter=12,
    )
    cell_st = ParagraphStyle(
        "Cell", parent=body_st, fontSize=9, leading=12, spaceAfter=0,
        alignment=0,
    )
    cell_hd = ParagraphStyle(
        "CellHd", parent=cell_st, fontName="Helvetica-Bold",
        textColor=colors.white,
    )
    title_st = ParagraphStyle(
        "Title2", parent=body_st, fontName="Helvetica-Bold", fontSize=21,
        leading=27, textColor=accent, alignment=TA_CENTER, spaceAfter=14,
    )
    sub_st = ParagraphStyle(
        "Sub", parent=body_st, fontName="Helvetica-Oblique", fontSize=12,
        leading=17, textColor=muted, alignment=TA_CENTER, spaceAfter=34,
    )
    meta_st = ParagraphStyle(
        "Meta", parent=body_st, fontSize=10.5, alignment=TA_CENTER, spaceAfter=5,
    )
    meta_mut = ParagraphStyle("MetaM", parent=meta_st, textColor=muted)
    disc_st = ParagraphStyle(
        "Disc", parent=body_st, fontName="Helvetica-Oblique", fontSize=8.5,
        leading=12, textColor=muted, alignment=TA_CENTER,
    )
    toc_st = ParagraphStyle(
        "Toc", parent=body_st, fontSize=10.5, leading=15, spaceAfter=6,
        leftIndent=14, alignment=0,
    )

    def footer(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(muted)
        canvas.drawCentredString(letter[0] / 2, 0.62 * inch, str(doc_.page))
        canvas.restoreState()

    def blank(canvas, doc_):
        pass

    story = []

    # Title page
    story += [Spacer(1, 2.5 * inch),
              Paragraph(TITLE, title_st),
              Paragraph(SUBTITLE, sub_st),
              Paragraph(DOCTYPE, meta_st),
              Paragraph(VERSION, meta_mut),
              Paragraph(DATE, meta_mut),
              Spacer(1, 1.6 * inch),
              Paragraph(DISCLAIMER, disc_st),
              NextPageTemplate("body"),
              PageBreak()]

    # Contents
    story.append(Paragraph("Contents", h1_st))
    for entry in toc_entries():
        story.append(Paragraph(entry, toc_st))
    story.append(PageBreak())

    for block in CONTENT:
        kind = block[0]
        if kind == "h1":
            story.append(Paragraph(block[1], h1_st))
        elif kind == "h2":
            story.append(Paragraph(block[1], h2_st))
        elif kind == "p":
            story.append(Paragraph(block[1], body_st))
        elif kind == "bullet":
            story.append(ListFlowable(
                [ListItem(Paragraph(block[1], body_st), leftIndent=18)],
                bulletType="bullet", start="•", leftIndent=16,
                bulletFontSize=9, spaceAfter=4,
            ))
        elif kind == "caption":
            story.append(Paragraph(block[1], cap_st))
        elif kind == "pagebreak":
            story.append(PageBreak())
        elif kind == "code":
            story.append(Preformatted("\n".join(block[1]), code_st))
        elif kind == "table":
            headers, rows, fracs = block[1], block[2], block[3]
            widths = [content_w * f for f in fracs]
            data = [[Paragraph(h, cell_hd) for h in headers]]
            data += [[Paragraph(c, cell_st) for c in row] for row in rows]
            t = Table(data, colWidths=widths, repeatRows=1)
            style = [
                ("BACKGROUND", (0, 0), (-1, 0), accent),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#C8CDD4")),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]
            for ri in range(1, len(data)):
                if ri % 2 == 0:
                    style.append(
                        ("BACKGROUND", (0, ri), (-1, ri), colors.HexColor("#" + ZEBRA_BG))
                    )
            t.setStyle(TableStyle(style))
            story += [t, Spacer(1, 6)]
        else:
            raise ValueError(f"unknown block: {kind}")

    doc = BaseDocTemplate(
        str(path), pagesize=letter,
        leftMargin=inch, rightMargin=inch, topMargin=inch, bottomMargin=inch,
        title=TITLE, author="Multi-Agent Medical Assessment", subject=SUBTITLE,
    )
    frame = Frame(inch, inch, content_w, letter[1] - 2 * inch, id="f")
    doc.addPageTemplates([
        PageTemplate(id="title", frames=[frame], onPage=blank),
        PageTemplate(id="body", frames=[frame], onPage=footer),
    ])
    doc.build(story)
    print(f"wrote {path}")


if __name__ == "__main__":
    out = Path(sys.argv[1])
    out.mkdir(parents=True, exist_ok=True)
    build_docx(out / "System-Documentation.docx")
    build_pdf(out / "System-Documentation.pdf")
